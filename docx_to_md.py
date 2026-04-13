#!/usr/bin/env python3
"""
DOCX to Markdown Exam Converter
Converts Word document exams into the markdown format required by md_to_jenzabar.py.

Supports:
- Cengage test bank exports (table-based with nested tables for choices/answers)
- Standard paragraph-based exam formats (numbered questions with lettered choices)

Usage:
    python docx_to_md.py exam.docx
    python docx_to_md.py exam.docx output_folder/

Requirements:
    pip install python-docx
"""

import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required.")
    print("Install it with: pip install python-docx")
    sys.exit(1)


# Metadata fields to strip (only keep ANSWER and RATIONALE)
STRIP_FIELDS = {
    'POINTS', 'QUESTION TYPE', 'HAS VARIABLES', 'LEARNING OBJECTIVES',
    'ACCREDITING STANDARDS', 'TOPICS', 'KEYWORDS', 'DATE CREATED',
    'DATE MODIFIED', 'DIFFICULTY', 'REFERENCES', 'OTHER', 'NOTES',
    'SECTION', 'CHAPTER', 'PAGE', 'OBJECTIVE', 'SKILL', 'LEVEL',
}


def parse_cengage_table_format(doc):
    """Parse Cengage test bank DOCX format (table-based with nested tables)."""
    questions = []

    for ti, table in enumerate(doc.tables):
        # Each top-level table is one question
        if len(table.rows) == 0 or len(table.columns) == 0:
            continue

        cell = table.cell(0, 0)

        # Get question text from cell paragraphs
        question_text = ''
        for p in cell.paragraphs:
            text = p.text.strip()
            if text:
                # Clean up non-breaking spaces
                text = text.replace('\xa0', ' ')
                question_text += text + ' '

        question_text = question_text.strip()
        if not question_text:
            continue

        # Extract question number and clean text
        q_match = re.match(r'^(\d+)\.\s*(.+)$', question_text)
        if q_match:
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()
        else:
            q_num = ti + 1
            q_text = question_text

        # Parse nested tables for choices and metadata
        choices = []
        answer = None
        rationale = None
        q_type_from_doc = None

        if cell.tables:
            for nested_table in cell.tables:
                for row in nested_table.rows:
                    cells_text = []
                    for c in row.cells:
                        cells_text.append(c.text.strip().replace('\xa0', ' '))

                    # Detect choice rows (have letter like a. b. c. in column 1)
                    if len(cells_text) >= 3:
                        letter_match = re.match(r'^([a-eA-E])\.?$', cells_text[1].strip())
                        if letter_match and cells_text[2].strip():
                            choices.append({
                                'letter': letter_match.group(1).upper(),
                                'text': cells_text[2].strip()
                            })
                            continue

                    # Detect metadata rows (label in column 0, value in column 1)
                    if len(cells_text) >= 2:
                        label = cells_text[0].strip().rstrip(':').upper().replace('\xa0', ' ')
                        value = cells_text[1].strip()

                        if label == 'ANSWER':
                            answer = value.upper()
                        elif label == 'RATIONALE':
                            rationale = value
                        elif label == 'QUESTION TYPE':
                            q_type_from_doc = value
                        # Skip all other metadata fields (POINTS, KEYWORDS, etc.)

        if not answer:
            continue

        questions.append({
            'num': q_num,
            'text': q_text,
            'choices': choices,
            'answer': answer,
            'rationale': rationale or '',
            'doc_type': q_type_from_doc
        })

    return questions


def parse_paragraph_format(doc):
    """Parse standard paragraph-based exam format."""
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip().replace('\xa0', ' ')
        if text:
            lines.append(text)

    if not lines:
        return []

    questions = []
    current_question = None
    current_num = 0
    current_choices = []
    current_answer = None
    current_rationale = None

    q_patterns = [
        re.compile(r'^\*?\*?(\d+)[\.\)]\s*(.+?)[\*]*$'),
        re.compile(r'^\((\d+)\)\s*(.+)$'),
        re.compile(r'^Question\s+(\d+)[:\.\)]\s*(.+)$', re.IGNORECASE),
    ]

    choice_patterns = [
        re.compile(r'^([A-Ea-e])\.\s+(.+)$'),
        re.compile(r'^([A-Ea-e])\)\s+(.+)$'),
        re.compile(r'^\(([A-Ea-e])\)\s+(.+)$'),
    ]

    answer_patterns = [
        re.compile(r'^\*?\*?Answer:\s*([A-Ea-e,\s]+|TRUE|FALSE)\*?\*?$', re.IGNORECASE),
        re.compile(r'^Correct Answer:\s*([A-Ea-e,\s]+|TRUE|FALSE)$', re.IGNORECASE),
        re.compile(r'^ANSWER:\s*([A-Ea-e,\s]+|TRUE|FALSE)$', re.IGNORECASE),
    ]

    rationale_patterns = [
        re.compile(r'^\*?\*?Rationale:\*?\*?\s*(.+)$', re.IGNORECASE),
        re.compile(r'^Explanation:\s*(.+)$', re.IGNORECASE),
    ]

    def save_question():
        nonlocal current_question, current_num, current_choices, current_answer, current_rationale
        if current_question and current_answer:
            questions.append({
                'num': current_num,
                'text': current_question,
                'choices': current_choices,
                'answer': current_answer.upper(),
                'rationale': current_rationale or '',
                'doc_type': None
            })
        current_question = None
        current_choices = []
        current_answer = None
        current_rationale = None

    for line in lines:
        # Check for question
        q_match = None
        for pat in q_patterns:
            q_match = pat.match(line)
            if q_match:
                break
        if q_match:
            save_question()
            current_num = int(q_match.group(1))
            current_question = q_match.group(2).strip().rstrip('*')
            continue

        # Check for choice
        choice_match = None
        for pat in choice_patterns:
            choice_match = pat.match(line)
            if choice_match:
                break
        if choice_match:
            current_choices.append({
                'letter': choice_match.group(1).upper(),
                'text': choice_match.group(2).strip()
            })
            continue

        # Check for answer
        answer_match = None
        for pat in answer_patterns:
            answer_match = pat.match(line)
            if answer_match:
                break
        if answer_match:
            current_answer = answer_match.group(1).strip()
            continue

        # Check for rationale
        rat_match = None
        for pat in rationale_patterns:
            rat_match = pat.match(line)
            if rat_match:
                break
        if rat_match:
            current_rationale = rat_match.group(1).strip()
            continue

    save_question()
    return questions


def detect_format(doc):
    """Detect if this is a table-based (Cengage) or paragraph-based format."""
    # If there are tables with nested tables, it is likely Cengage format
    if doc.tables:
        cell = doc.tables[0].cell(0, 0) if doc.tables[0].rows else None
        if cell and cell.tables:
            return 'cengage'

    # Check for paragraph content
    has_paragraphs = any(p.text.strip() for p in doc.paragraphs)
    if has_paragraphs:
        return 'paragraph'

    # Fallback: check if tables have question-like content
    if doc.tables:
        return 'cengage'

    return 'unknown'


def questions_to_markdown(questions, title):
    """Convert parsed questions to the standard markdown format."""
    md_lines = [f"# {title}", ""]

    for q in questions:
        md_lines.append(f"**{q['num']}. {q['text']}**")
        md_lines.append("")

        answer = q['answer'].strip()

        if answer in ('TRUE', 'FALSE'):
            # True/False question, no choices
            pass
        else:
            for choice in q['choices']:
                md_lines.append(f"{choice['letter']}. {choice['text']}")
            md_lines.append("")

        md_lines.append(f"**Answer: {answer}**")
        md_lines.append("")

        if q['rationale']:
            md_lines.append(f"**Rationale:** {q['rationale']}")
            md_lines.append("")

        md_lines.append("---")
        md_lines.append("")

    return '\n'.join(md_lines)


def convert_docx_to_md(docx_path, output_dir=None):
    """Convert a DOCX exam file to markdown format."""
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}")
        return None

    print(f"Reading: {docx_path}")
    doc = Document(docx_path)

    # Detect format
    fmt = detect_format(doc)
    print(f"  Detected format: {fmt}")

    if fmt == 'cengage':
        questions = parse_cengage_table_format(doc)
    elif fmt == 'paragraph':
        questions = parse_paragraph_format(doc)
    else:
        print("  WARNING: Could not detect exam format.")
        # Try both and use whichever finds more questions
        q1 = parse_cengage_table_format(doc)
        q2 = parse_paragraph_format(doc)
        questions = q1 if len(q1) >= len(q2) else q2

    print(f"  Parsed {len(questions)} questions")

    if not questions:
        print("  WARNING: No questions found.")
        return None

    # Count types
    tf_count = sum(1 for q in questions if q['answer'].upper() in ('TRUE', 'FALSE'))
    multi_count = sum(1 for q in questions if ',' in q['answer'])
    mc_count = len(questions) - tf_count - multi_count
    print(f"  - Multiple Choice: {mc_count}")
    print(f"  - Multiple Answer: {multi_count}")
    print(f"  - True/False: {tf_count}")

    # Title from filename
    title = os.path.splitext(os.path.basename(docx_path))[0]
    title = title.replace('_', ' ')

    md_content = questions_to_markdown(questions, title)

    # Output path
    if output_dir is None:
        output_dir = os.path.dirname(docx_path) or '.'

    os.makedirs(output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    md_path = os.path.join(output_dir, f"{base_name}.md")

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    print(f"  Saved: {md_path}")
    return md_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python docx_to_md.py <exam.docx> [output_folder]")
        print("")
        print("Converts a Word document exam to markdown format.")
        print("Supports Cengage test bank exports and standard exam formats.")
        print("")
        print("The markdown file can then be converted to a Jenzabar cartridge with:")
        print("  python md_to_jenzabar.py <exam.md>")
        sys.exit(1)

    docx_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    result = convert_docx_to_md(docx_file, output_dir)
    if result:
        print(f"\nDone. Now convert to Jenzabar cartridge with:")
        print(f"  python md_to_jenzabar.py \"{result}\"")
